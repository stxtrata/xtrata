from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Mm, Pt, RGBColor


OUT = Path("output/docx/Jim_Crane_CV-2026.docx")

BLACK = RGBColor(0, 0, 0)
MUTED = RGBColor(78, 78, 78)
ACCENT = RGBColor(42, 92, 122)


def set_font(run, size=9.5, bold=False, italic=False, color=BLACK):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def set_para_spacing(paragraph, before=0, after=3, line=1.04):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_rule(paragraph, color="DADCE0", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_p(doc, text="", size=9.5, bold=False, italic=False, color=BLACK, after=3, before=0, align=None, line=1.04):
    p = doc.add_paragraph()
    set_para_spacing(p, before=before, after=after, line=line)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def add_hyperlink(paragraph, text, url, size=9.1, color=ACCENT):
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = paragraph.add_run(text)
    set_font(run, size=size, color=color)
    run.font.underline = True
    hyperlink.append(run._r)
    paragraph._p.append(hyperlink)


def add_link_paragraph(doc, segments, size=9.1, color=BLACK, after=1.5, align=None):
    p = doc.add_paragraph()
    set_para_spacing(p, after=after, line=1.04)
    if align is not None:
        p.alignment = align
    for text, url in segments:
        if url:
            add_hyperlink(p, text, url, size=size)
        else:
            run = p.add_run(text)
            set_font(run, size=size, color=color)
    return p


def add_heading(doc, text, size=11.4, before=8, after=2):
    p = add_p(doc, text.upper(), size=size, bold=True, color=ACCENT, before=before, after=2)
    return p


def add_bullet(doc, text, size=9.5, after=1.6):
    p = doc.add_paragraph(style="List Bullet")
    set_para_spacing(p, before=0, after=after, line=1.02)
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    r = p.add_run(text)
    set_font(r, size=size)
    return p


def add_role(doc, title, org, dates):
    p = doc.add_paragraph()
    set_para_spacing(p, before=4, after=0.7, line=1.02)
    r = p.add_run(title)
    set_font(r, size=9.9, bold=True)
    r = p.add_run(f" | {org}")
    set_font(r, size=9.9, color=MUTED)
    r = p.add_run(f" | {dates}")
    set_font(r, size=9.9, color=MUTED)
    return p


def add_two_col_keywords(doc, left, right):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    widths = [Inches(3.1), Inches(3.1)]
    for i, cell in enumerate(table.rows[0].cells):
        cell.width = widths[i]
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = tc_pr.find(qn("w:tcW"))
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
            tc_pr.append(tc_w)
        tc_w.set(qn("w:w"), str(int(widths[i].inches * 1440)))
        tc_w.set(qn("w:type"), "dxa")

    for cell, items in zip(table.rows[0].cells, [left, right]):
        cell.text = ""
        for item in items:
            p = cell.add_paragraph(style="List Bullet")
            set_para_spacing(p, after=1.2, line=1.0)
            p.paragraph_format.left_indent = Inches(0.17)
            p.paragraph_format.first_line_indent = Inches(-0.12)
            r = p.add_run(item)
            set_font(r, size=8.9)
    return table


def set_cell_margins(table, top=40, start=40, bottom=40, end=40):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tc_pr = tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
                node = tc_mar.find(qn(f"w:{m}"))
                if node is None:
                    node = OxmlElement(f"w:{m}")
                    tc_mar.append(node)
                node.set(qn("w:w"), str(v))
                node.set(qn("w:type"), "dxa")


def configure_section(section):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)


def build():
    doc = Document()
    section = doc.sections[0]
    configure_section(section)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)
    styles["List Bullet"].font.name = "Arial"
    styles["List Bullet"].font.size = Pt(9.5)

    p = add_p(doc, "Jim Crane", size=22, bold=True, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = add_p(
        doc,
        "On-Chain Artist | Creative Technologist | Sound Systems & Digital Archives",
        size=10.8,
        color=MUTED,
        after=2,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    p = add_p(
        doc,
        "Brighton, UK | +44 7765 241717 | jimdotbtc@gmail.com",
        size=8.6,
        color=MUTED,
        after=1,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    p = add_link_paragraph(
        doc,
        [
            ("xtrata.xyz", "https://xtrata.xyz"),
            (" | ", None),
            ("audionals.com", "https://audionals.com"),
            (" | ", None),
            ("x.com/jimdotbtc", "https://x.com/jimdotbtc"),
        ],
        size=8.6,
        color=MUTED,
        after=4,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_rule(p)

    add_heading(doc, "Profile", before=5)
    add_p(
        doc,
        "Artist, creative technologist and sound engineer building permanent, verifiable systems for on-chain art and music. My practice uses blockchain as both medium and archive: origin material such as loops, stems, credits, metadata, code, interviews and project histories becomes part of the work rather than hidden production residue. Since 2014, I have developed Bitcoin and Stacks projects spanning recursive media, provenance, rights, preservation, public research archives and AI-assisted creative workflows.",
        size=9.6,
        after=4,
        line=1.06,
    )

    add_heading(doc, "Practice Focus", before=6, after=1)
    add_p(
        doc,
        "On-chain creation | Provenance and source material | Creative archives and cultural memory | Rights and re-use | Human and machine authorship | Generative audio and sound installation",
        size=9.4,
        color=MUTED,
        after=4,
    )

    add_heading(doc, "Selected Onchain & Creative Technology Work", before=7)

    add_role(doc, "Founder and Developer", "Xtrata / Forever Twins", "2026-present")
    add_bullet(doc, "Created Xtrata, an inscription-based data layer on Stacks for permanent, fully on-chain generative art and programmable cultural records.")
    add_bullet(doc, "Created Forever Twins, giving existing NFT collections permanent, hash-verified twins inscribed on Bitcoin via Stacks so their artwork survives failed off-chain links.")
    add_bullet(doc, "Awarded a Stacks Community DeGrants Cohort 4 grant; launched the first public collection, Bitcoin Pepes, with holders already claiming twins.")
    add_bullet(doc, "Built XTRATA FM, an on-chain radio widget that streams multi-megabyte tracks directly from inscriptions.")

    add_role(doc, "Founder", "Audionals", "2023-present")
    add_bullet(doc, "Developed a protocol for producing, storing and rights-managing music directly on Bitcoin, treating stems, credits, provenance and reusable components as native parts of the work.")
    add_bullet(doc, "Released TRUTH, the first recursive music collection on Bitcoin; the collection sold out in just over an hour.")
    add_bullet(doc, "Used recursive references to reduce the incremental inscription data for tracks derived from roughly 70 MB of music to around 3 KB per track, making fully on-chain music economically viable within Bitcoin's block-size limits.")
    add_bullet(doc, "Developing KP Loops with Kieron Pepper (The Prodigy), an on-chain drum-loop library designed around traceable source material and re-use.")

    add_role(doc, "Creator, Host and Producer", "The Audionals Show", "2023-present")
    add_bullet(doc, "Created a long-form interview and discussion series on on-chain music, sound design and creative technology within the Bitcoin ecosystem.")
    add_bullet(doc, "Built a public research archive through 40+ episodes, leading concept development, guest curation, technical production and documentation.")

    add_role(doc, "Co-founder", "This Is Number One", "2021")
    add_bullet(doc, "Co-founded the first NFT marketplace on the Stacks blockchain, debuting the network's first commercial and music NFT collections.")
    add_bullet(doc, "Collaborators included Fatboy Slim, Orbital, Cara Delevingne, Dave Stewart (Eurythmics) and Chemical X.")

    add_role(doc, "Creative Technologist", "AI and Audio Workflow Systems", "2024-present")
    add_bullet(doc, "Designed automation tools that reduced a publisher's audiobook production workflow from roughly 8 hours to 6 minutes, translating unstructured text and media tasks into repeatable pipelines.")
    add_bullet(doc, "Use AI systems for research, prototyping, editing, compression and media structuring, with attention to authorship, traceability and human judgement.")

    doc.add_section(WD_SECTION.NEW_PAGE)
    section = doc.sections[-1]
    configure_section(section)

    add_heading(doc, "Selected Installation & Sound Design", before=0)
    add_role(doc, "Sound and Lighting Designer / Technician", "THEY / ONLAR by Ipek Duben, Fabrica Gallery", "2017")
    add_bullet(doc, "Designed and deployed directional speaker systems, multi-zone playback and sensor-driven interactivity for an immersive sound installation.")

    add_role(doc, "Audio System Designer and Lead PA Engineer", "Stories in Motion / Longrange, Brighton Festival", "2006")
    add_bullet(doc, "Designed and implemented a quadraphonic PA configuration for a performance featuring Irvine Welsh and Longrange.")

    add_heading(doc, "Professional Technical Background", before=8)

    add_role(doc, "Sound Engineer and Technician", "Brighton Dome & Festival", "2002-present")
    add_bullet(doc, "Deliver sound, stage and technical support across music, theatre, dance, comedy, festivals and outdoor arts, realising complex work with visiting artists and production companies.")

    add_role(doc, "Tour Manager, Front of House and Monitor Engineer", "Freelance", "2006-2019")
    add_bullet(doc, "Managed international tours and led crews of up to 15 for artists including Palma Violets, The Go! Team, The Big Moon, Kelis and Tom Waits.")

    add_role(doc, "Monitor Engineer", "Krankenhaus Festival, Muncaster", "2025")
    add_bullet(doc, "Provided monitor engineering and stage sound management for the multi-day independent arts and music festival founded by Sea Power.")

    add_role(doc, "Additional Selected Production Credits", "BBC, ITV, Glastonbury Festival, Brighton Festival", "2002-2012")
    add_bullet(doc, "Moving light technician for Later... with Jools Holland and The Jonathan Ross Show; stage crew on Glastonbury's Jazz World / One World stage; audio system designer and lead PA engineer for a quadraphonic Brighton Festival project featuring Irvine Welsh and Longrange.")

    add_heading(doc, "Education and Training", before=8)
    add_bullet(doc, "Diploma in Audio Engineering - School of Audio Engineers, Sydney, Australia")
    add_bullet(doc, "Diploma in Performing Arts - The Actors Centre, Sydney, Australia")
    add_bullet(doc, "Ongoing self-directed study in blockchain development, AI-assisted production, generative audio, inscriptions, creative metadata and decentralised cultural infrastructure.")

    add_heading(doc, "Selected Milestones", before=8)
    add_bullet(doc, "Imported the UK's first two-way Bitcoin ATM in 2014, beginning a long-term engagement with decentralised systems.")
    add_bullet(doc, "Helped pioneer the first NFTs on Stacks in 2021 and establish early models for Bitcoin-adjacent creative publishing.")

    add_heading(doc, "Links", before=8)
    add_link_paragraph(
        doc,
        [
            ("Xtrata: ", None),
            ("https://xtrata.xyz", "https://xtrata.xyz"),
            (" | Audionals: ", None),
            ("https://audionals.com", "https://audionals.com"),
        ],
        size=9.1,
    )
    add_link_paragraph(
        doc,
        [
            ("KP Loops: ", None),
            ("https://audionals.com/BAM/kp-loops", "https://audionals.com/BAM/kp-loops"),
            (" | X / Twitter: ", None),
            ("https://x.com/jimdotbtc", "https://x.com/jimdotbtc"),
        ],
        size=9.1,
    )

    for section in doc.sections:
        section.footer.is_linked_to_previous = False
        footer = section.footer.paragraphs[0]
        footer.clear()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(footer, after=0)
        r = footer.add_run("Jim Crane - CV for Padimai OOPS Onchain Artist Residency")
        set_font(r, size=7.5, color=MUTED)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
