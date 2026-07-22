#!/usr/bin/env python3
"""Professional styling pass for md2docx output (footnote route).

Grayscale-only palette, Georgia body / Calibri headings, generous margins,
title page block, light-formatted tables, hanging-indent references,
footer page numbers. No content changes.
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = '/mnt/agents/output/docx_build/report.footnote.docx'
DST = '/mnt/agents/output/Xtrata_Use_Cases_Report.docx'

INK      = RGBColor(0x1A, 0x1A, 0x1A)   # near-black body
H2INK    = RGBColor(0x1F, 0x1F, 0x1F)
H3INK    = RGBColor(0x26, 0x26, 0x26)
H4INK    = RGBColor(0x33, 0x33, 0x33)
GRAY     = RGBColor(0x40, 0x40, 0x40)   # footnotes / subtitle
LTGRAY   = RGBColor(0x59, 0x59, 0x59)   # date / page numbers
RULE     = 'BFBFBF'                      # light border
RULE_LT  = 'D9D9D9'
HDR_FILL = 'F2F2F2'                      # table header shading

SERIF = 'Georgia'
SANS  = 'Calibri'
MONO  = 'Consolas'

doc = Document(SRC)


def set_style_font(st, name=None, size=None, bold=None, italic=None, color=None,
                   small_caps=None):
    f = st.font
    if name is not None:
        f.name = name
    if size is not None:
        f.size = Pt(size)
    if bold is not None:
        f.bold = bold
    if italic is not None:
        f.italic = italic
    if color is not None:
        f.color.rgb = color
    if small_caps is not None:
        f.small_caps = small_caps


def set_pf(st, before=None, after=None, line=None, keep_next=None):
    pf = st.paragraph_format
    if before is not None:
        pf.space_before = Pt(before)
    if after is not None:
        pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line
    if keep_next is not None:
        pf.keep_with_next = keep_next


def style_bottom_border(st, color, sz='4'):
    pPr = st.element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


styles = {s.style_id: s for s in doc.styles}

# ── Base text styles ──────────────────────────────────────────────
set_style_font(styles['Normal'], SERIF, 10.5, color=INK)
set_pf(styles['Normal'], after=7, line=1.15)
styles['Normal'].paragraph_format.widow_control = True

set_style_font(styles['BodyText'], SERIF, 10.5, color=INK)
set_pf(styles['BodyText'], after=7, line=1.15)
set_style_font(styles['FirstParagraph'], SERIF, 10.5, color=INK)
set_pf(styles['FirstParagraph'], after=7, line=1.15)

# Table cell text
set_style_font(styles['Compact'], SANS, 10, color=H3INK)
set_pf(styles['Compact'], after=2, line=1.05)

# ── Heading hierarchy (dark gray/black only) ─────────────────────
set_style_font(styles['Heading1'], SANS, 17, bold=True, italic=False, color=INK)
set_pf(styles['Heading1'], before=0, after=12, keep_next=True)
style_bottom_border(styles['Heading1'], RULE, sz='6')

set_style_font(styles['Heading2'], SANS, 15, bold=True, italic=False, color=H2INK)
set_pf(styles['Heading2'], before=22, after=8, keep_next=True)
style_bottom_border(styles['Heading2'], RULE_LT, sz='4')

set_style_font(styles['Heading3'], SANS, 12.5, bold=True, italic=False, color=H3INK)
set_pf(styles['Heading3'], before=14, after=5, keep_next=True)

set_style_font(styles['Heading4'], SANS, 11, bold=True, italic=False, color=H4INK)
set_pf(styles['Heading4'], before=10, after=4, keep_next=True)

# ── Title page styles ─────────────────────────────────────────────
set_style_font(styles['Title'], SANS, 25, bold=True, color=INK)
set_pf(styles['Title'], before=0, after=16, line=1.1)

set_style_font(styles['Subtitle'], SANS, 13, bold=False, italic=True, color=GRAY)
set_pf(styles['Subtitle'], before=0, after=6, line=1.2)

set_style_font(styles['Date'], SANS, 12, bold=False, italic=False, color=LTGRAY,
               small_caps=True)
set_pf(styles['Date'], before=10, after=0)

# ── Footnotes / character styles ─────────────────────────────────
set_style_font(styles['FootnoteText'], SERIF, 8.5, color=GRAY)
set_pf(styles['FootnoteText'], after=2, line=1.0)
set_style_font(styles['FootnoteReference'], color=INK)
set_style_font(styles['VerbatimChar'], MONO, 9, color=H4INK)
set_style_font(styles['Hyperlink'], color=INK)

# ── Page setup: generous margins ─────────────────────────────────
sec = doc.sections[0]
sec.top_margin = Inches(1.1)
sec.bottom_margin = Inches(1.0)
sec.left_margin = Inches(1.25)
sec.right_margin = Inches(1.25)
sec.different_first_page_header_footer = True

# ── Footer: centered page number (suppressed on title page) ──────
footer = sec.footer
fp = footer.paragraphs[0]
for r in list(fp.runs):
    r._element.getparent().remove(r._element)
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run()
fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
instr.text = 'PAGE'
fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
run.font.name = SERIF
run.font.size = Pt(9)
run.font.color.rgb = LTGRAY

# ── Title block restructure ───────────────────────────────────────
paras = doc.paragraphs
title_p = paras[0]
assert title_p.style.name == 'Heading 1', title_p.style.name
title_p.style = doc.styles['Title']
title_p.paragraph_format.space_before = Pt(110)
# rule under the title
tpPr = title_p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
b = OxmlElement('w:bottom')
b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '8')
b.set(qn('w:space'), '10'); b.set(qn('w:color'), '737373')
pBdr.append(b); tpPr.append(pBdr)

subtitle_p = paras[1]
subtitle_p.style = doc.styles['Subtitle']
subtitle_p.runs[0].text = 'A use-case and launch-strategy report based on deep research'

# Date line + page break after it (body starts on page 2)
exec_h = paras[2]
assert exec_h.style.name == 'Heading 2' and exec_h.text.strip() == 'Executive Summary'
date_p = exec_h.insert_paragraph_before('July 21, 2026', style=doc.styles['Date'])
date_p.add_run().add_break(WD_BREAK.PAGE)

# References chapter starts on its own page
for p in doc.paragraphs:
    if p.text.strip() == 'References' and p.style.name == 'Heading 1':
        p.paragraph_format.page_break_before = True
        break

# ── References entries: hanging indent, compact ──────────────────
in_refs = False
ref_count = 0
for p in doc.paragraphs:
    if p.text.strip() == 'References':
        in_refs = True
        continue
    if in_refs and re.match(r'^\[\d+\]', p.text.strip()):
        pf = p.paragraph_format
        pf.left_indent = Inches(0.35)
        pf.first_line_indent = Inches(-0.35)
        pf.space_after = Pt(5)
        pf.line_spacing = 1.08
        for r in p.runs:
            r.font.name = SERIF
            r.font.size = Pt(9.5)
            r.font.color.rgb = H3INK
        ref_count += 1

# ── Tables: light borders, shaded header, repeat header row ──────
def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

for tbl in doc.tables:
    tblPr = tbl._tbl.tblPr
    # borders
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), RULE)
        borders.append(e)
    tblPr.append(borders)
    # cell padding
    for old in tblPr.findall(qn('w:tblCellMar')):
        tblPr.remove(old)
    mar = OxmlElement('w:tblCellMar')
    for side, w in (('top', '50'), ('left', '100'), ('bottom', '50'), ('right', '100')):
        e = OxmlElement(f'w:{side}')
        e.set(qn('w:w'), w); e.set(qn('w:type'), 'dxa')
        mar.append(e)
    tblPr.append(mar)
    # header row
    hdr = tbl.rows[0]
    trPr = hdr._tr.get_or_add_trPr()
    if trPr.find(qn('w:tblHeader')) is None:
        th = OxmlElement('w:tblHeader'); th.set(qn('w:val'), 'true')
        trPr.append(th)
    for cell in hdr.cells:
        shade(cell, HDR_FILL)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = INK

# ── Core properties ───────────────────────────────────────────────
cp = doc.core_properties
cp.title = 'Xtrata: The Permanent Object Layer — Revolutionary Use Cases'
cp.subject = 'Use-case and launch-strategy report'
cp.author = 'Deep Research'
cp.comments = 'Generated from Markdown via md2docx (footnote citation style)'

doc.save(DST)
print(f'Saved {DST} | reference paragraphs styled: {ref_count}')
